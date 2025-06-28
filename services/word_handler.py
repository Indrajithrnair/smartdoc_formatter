from docx import Document
import os

class WordHandler:
    @staticmethod
    def create_document():
        """
        Create a new Word document
        """
        return Document()

    @staticmethod
    def read_document(file_path):
        """
        Read a Word document and extract its content
        """
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return {
                'status': 'success',
                'content': content
            }
        except Exception as e:
            return {
                'status': 'error',
                'message': f'Error reading document: {str(e)}'
            }

    @staticmethod
    def save_document(document, output_path):
        """
        Save the Word document to the specified path
        """
        try:
            document.save(output_path)
            return {
                'status': 'success',
                'message': 'Document saved successfully'
            }
        except Exception as e:
            return {
                'status': 'error',
                'message': f'Error saving document: {str(e)}'
            } 